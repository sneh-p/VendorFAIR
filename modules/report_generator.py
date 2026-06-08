"""Risk memo generation — DOCX (python-docx) and PDF (reportlab)."""
import io
import json
import logging
from datetime import date
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

import config
from modules.fair_calculator import FairResults
from prompts.report_writer import (
    CONTROLS_PROMPT,
    EXEC_SUMMARY_PROMPT,
    FALLBACK_CONTROLS,
    RECOMMENDATION_BY_TIER,
)

logger = logging.getLogger(__name__)


def _safe_name(name: str) -> str:
    return "".join(c for c in name if c.isalnum() or c in ("-", "_")) or "Unknown"


def report_filename(tenant_name: str, vendor_name: str, ext: str) -> str:
    return (
        f"VendorFAIR_{_safe_name(tenant_name)}_{_safe_name(vendor_name)}_"
        f"{date.today().isoformat()}.{ext}"
    )


def render_risk_chart(results: FairResults) -> bytes:
    """Render the ALE histogram with percentile markers as PNG bytes."""
    fig, ax = plt.subplots(figsize=(7, 3.5), dpi=150)
    ax.hist(results.ale_samples, bins=60, color="#3498db", alpha=0.8, edgecolor="white")
    for label, value, color in (
        ("P10", results.p10, "#2ecc71"),
        ("P50", results.p50, "#f39c12"),
        ("P90", results.p90, "#e74c3c"),
    ):
        ax.axvline(value, color=color, linestyle="--", linewidth=1.5,
                   label=f"{label}: ${value:,.0f}")
    ax.set_xlabel("Annualized Loss Exposure ($)")
    ax.set_ylabel("Frequency")
    ax.set_title("FAIR Monte Carlo — Simulated Annual Loss Distribution")
    ax.legend()
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return buf.getvalue()


def _ai_text(prompt: str) -> str | None:
    """Call Claude for a report section; return None when unavailable."""
    if not config.ANTHROPIC_API_KEY:
        return None
    try:
        import anthropic

        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
        response = client.messages.create(
            model=config.CLAUDE_MODEL,
            max_tokens=1024,
            messages=[{"role": "user", "content": prompt}],
        )
        return next((b.text for b in response.content if b.type == "text"), None)
    except Exception:  # noqa: BLE001 — report generation must not fail on AI errors
        logger.exception("AI report-section generation failed; using fallback")
        return None


def build_report_sections(vendor, tenant, assessment, results: FairResults) -> dict:
    """Assemble all memo sections as plain data, used by both DOCX and PDF."""
    research = {}
    if assessment.research_json:
        try:
            research = json.loads(assessment.research_json)
        except json.JSONDecodeError:
            pass

    risk_tier = results.risk_tier
    research_summary = research.get("research_summary", "No AI research performed.")

    exec_summary = _ai_text(
        EXEC_SUMMARY_PROMPT.format(
            vendor_name=vendor.name,
            category=vendor.category,
            tenant_name=tenant.name,
            ale_p50=results.p50,
            risk_tier=risk_tier,
            public_trust_posture=research.get("public_trust_posture", "unknown"),
            research_summary=research_summary,
        )
    ) or (
        f"{vendor.name} presents a {risk_tier.lower()} level of annualized risk to "
        f"{tenant.name}, with an estimated median annual loss exposure of "
        f"${results.p50:,.0f}. The recommended treatment is: "
        f"{RECOMMENDATION_BY_TIER[risk_tier]}."
    )

    controls_text = _ai_text(
        CONTROLS_PROMPT.format(
            vendor_name=vendor.name,
            category=vendor.category,
            risk_tier=risk_tier,
            data_types=assessment.data_types_accessed,
            integration_depth=assessment.integration_depth,
            regulatory_context=assessment.regulatory_context,
            research_summary=research_summary,
        )
    )
    controls = (
        [line for line in controls_text.splitlines() if line.strip()]
        if controls_text
        else FALLBACK_CONTROLS[risk_tier]
    )

    return {
        "title": f"Vendor Risk Assessment — {vendor.name}",
        "exec_summary": exec_summary,
        "vendor_profile": [
            ("Vendor", vendor.name),
            ("Website", vendor.website or "—"),
            ("Category", vendor.category),
            ("Client Tenant", tenant.name),
            ("Contract Status", vendor.contract_status),
            ("Data Types Accessed", assessment.data_types_accessed or "—"),
            ("Integration Depth", assessment.integration_depth or "—"),
            ("Regulatory Context", assessment.regulatory_context or "—"),
            ("Assessor", assessment.assessor or "—"),
        ],
        "research": research,
        "fair": [
            ("ALE P10", f"${results.p10:,.0f}"),
            ("ALE P50 (median)", f"${results.p50:,.0f}"),
            ("ALE P90", f"${results.p90:,.0f}"),
            ("Risk Tier", risk_tier),
        ],
        "recommendation": RECOMMENDATION_BY_TIER[risk_tier],
        "risk_tier": risk_tier,
        "controls": controls,
        "fair_inputs": [
            ("TEF (events/yr)", assessment.tef_min, assessment.tef_ml, assessment.tef_max),
            ("TC (0-1)", assessment.tc_min, assessment.tc_ml, assessment.tc_max),
            ("CS (0-1)", assessment.cs_min, assessment.cs_ml, assessment.cs_max),
            ("PLM ($)", assessment.plm_min, assessment.plm_ml, assessment.plm_max),
            ("SLM ($)", assessment.slm_min, assessment.slm_ml, assessment.slm_max),
        ],
        "iterations": len(results.ale_samples),
    }


def generate_docx(vendor, tenant, assessment, results: FairResults, output_dir=None) -> str:
    """Generate the DOCX risk memo; returns the file path."""
    sections = build_report_sections(vendor, tenant, assessment, results)
    chart_png = render_risk_chart(results)

    doc = Document()
    doc.add_heading(sections["title"], level=0)
    doc.add_paragraph(f"{config.ORG_NAME} — Prepared {date.today().isoformat()}")

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(sections["exec_summary"])

    doc.add_heading("2. Vendor Profile", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, value in sections["vendor_profile"]:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = str(value)

    doc.add_heading("3. Research Findings", level=1)
    research = sections["research"]
    if research:
        doc.add_paragraph(research.get("research_summary", ""))
        certs = research.get("certifications_found") or []
        if certs:
            doc.add_paragraph("Certifications found: " + ", ".join(map(str, certs)))
        breaches = research.get("breach_history") or []
        if breaches:
            doc.add_paragraph("Breach history:")
            for b in breaches:
                doc.add_paragraph(str(b), style="List Bullet")
        links = research.get("evidence_links") or []
        if links:
            doc.add_paragraph("Evidence links:")
            for link in links:
                doc.add_paragraph(str(link), style="List Bullet")
        doc.add_paragraph(
            f"Research confidence: {research.get('research_confidence', 'n/a')}"
        )
    else:
        doc.add_paragraph("No AI research was performed for this assessment.")

    doc.add_heading("4. FAIR Risk Analysis", level=1)
    doc.add_picture(io.BytesIO(chart_png), width=Inches(6.0))
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, value in sections["fair"]:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = str(value)

    doc.add_heading("5. Risk Tier & Recommendation", level=1)
    doc.add_paragraph(
        f"Risk tier: {sections['risk_tier']}. Recommended treatment: "
        f"{sections['recommendation']}."
    )

    doc.add_heading("6. Required Controls / Contractual Clauses", level=1)
    for control in sections["controls"]:
        doc.add_paragraph(control.lstrip("- ").strip(), style="List Bullet")

    doc.add_heading("7. Reviewer Sign-off", level=1)
    doc.add_paragraph("Reviewed by: ______________________    Date: ____________")
    doc.add_paragraph("Approved by: ______________________    Date: ____________")

    doc.add_heading("8. Appendix: FAIR Inputs & Monte Carlo Parameters", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(("Parameter", "Min", "Most Likely", "Max")):
        hdr[i].text = h
    for name, lo, ml, hi in sections["fair_inputs"]:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"{lo:,.3g}"
        row[2].text = f"{ml:,.3g}"
        row[3].text = f"{hi:,.3g}"
    doc.add_paragraph(f"Monte Carlo iterations: {sections['iterations']:,}")

    out_dir = Path(output_dir or config.REPORT_OUTPUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / report_filename(tenant.name, vendor.name, "docx")
    doc.save(str(path))
    return str(path)


def generate_pdf(vendor, tenant, assessment, results: FairResults, output_dir=None) -> str:
    """Generate the PDF risk memo; returns the file path."""
    sections = build_report_sections(vendor, tenant, assessment, results)
    chart_png = render_risk_chart(results)

    out_dir = Path(output_dir or config.REPORT_OUTPUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / report_filename(tenant.name, vendor.name, "pdf")

    styles = getSampleStyleSheet()
    story = [
        Paragraph(sections["title"], styles["Title"]),
        Paragraph(f"{config.ORG_NAME} — Prepared {date.today().isoformat()}", styles["Normal"]),
        Spacer(1, 12),
        Paragraph("1. Executive Summary", styles["Heading1"]),
        Paragraph(sections["exec_summary"], styles["Normal"]),
        Spacer(1, 8),
        Paragraph("2. Vendor Profile", styles["Heading1"]),
    ]

    def data_table(rows):
        t = Table([[str(k), str(v)] for k, v in rows], colWidths=[2.2 * inch, 4.2 * inch])
        t.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BACKGROUND", (0, 0), (0, -1), colors.whitesmoke),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        return t

    story.append(data_table(sections["vendor_profile"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("3. Research Findings", styles["Heading1"]))
    research = sections["research"]
    if research:
        story.append(Paragraph(research.get("research_summary", ""), styles["Normal"]))
        certs = research.get("certifications_found") or []
        if certs:
            story.append(
                Paragraph("Certifications found: " + ", ".join(map(str, certs)), styles["Normal"])
            )
        for link in research.get("evidence_links") or []:
            story.append(Paragraph(f"• {link}", styles["Normal"]))
        story.append(
            Paragraph(
                f"Research confidence: {research.get('research_confidence', 'n/a')}",
                styles["Normal"],
            )
        )
    else:
        story.append(
            Paragraph("No AI research was performed for this assessment.", styles["Normal"])
        )
    story.append(Spacer(1, 8))

    story.append(Paragraph("4. FAIR Risk Analysis", styles["Heading1"]))
    story.append(RLImage(io.BytesIO(chart_png), width=6.2 * inch, height=3.1 * inch))
    story.append(data_table(sections["fair"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("5. Risk Tier & Recommendation", styles["Heading1"]))
    story.append(
        Paragraph(
            f"Risk tier: {sections['risk_tier']}. Recommended treatment: "
            f"{sections['recommendation']}.",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 8))

    story.append(Paragraph("6. Required Controls / Contractual Clauses", styles["Heading1"]))
    for control in sections["controls"]:
        story.append(Paragraph(f"• {control.lstrip('- ').strip()}", styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("7. Reviewer Sign-off", styles["Heading1"]))
    story.append(Paragraph("Reviewed by: ______________________  Date: ____________", styles["Normal"]))
    story.append(Paragraph("Approved by: ______________________  Date: ____________", styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("8. Appendix: FAIR Inputs & Monte Carlo Parameters", styles["Heading1"]))
    appendix_rows = [["Parameter", "Min", "Most Likely", "Max"]] + [
        [name, f"{lo:,.3g}", f"{ml:,.3g}", f"{hi:,.3g}"]
        for name, lo, ml, hi in sections["fair_inputs"]
    ]
    t = Table(appendix_rows, colWidths=[2.0 * inch, 1.5 * inch, 1.5 * inch, 1.5 * inch])
    t.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    story.append(t)
    story.append(Paragraph(f"Monte Carlo iterations: {sections['iterations']:,}", styles["Normal"]))

    SimpleDocTemplate(str(path), pagesize=letter).build(story)
    return str(path)
