"""Report generation tests — DOCX and PDF export against a temp database."""
import json
from datetime import datetime
from types import SimpleNamespace

import pytest
from docx import Document

from modules.fair_calculator import FairInputs, run_simulation
from modules.report_generator import (
    generate_docx,
    generate_pdf,
    render_risk_chart,
    report_filename,
)


@pytest.fixture
def fixtures():
    tenant = SimpleNamespace(name="Test Tenant")
    vendor = SimpleNamespace(
        name="Acme SaaS",
        website="https://acme.example",
        category="SaaS",
        contract_status="in evaluation",
        tenant=tenant,
    )
    assessment = SimpleNamespace(
        assessor="Tester",
        data_types_accessed="PII, financial",
        integration_depth="API/SSO",
        regulatory_context="SOC2, PIPEDA",
        research_json=json.dumps(
            {
                "research_summary": "Strong posture.",
                "certifications_found": ["SOC 2 Type II"],
                "breach_history": [],
                "evidence_links": ["https://acme.example/trust"],
                "research_confidence": "high",
                "public_trust_posture": "strong",
            }
        ),
        tef_min=0.1, tef_ml=0.5, tef_max=2.0,
        tc_min=0.3, tc_ml=0.6, tc_max=0.85,
        cs_min=0.4, cs_ml=0.65, cs_max=0.9,
        plm_min=10_000.0, plm_ml=75_000.0, plm_max=500_000.0,
        slm_min=5_000.0, slm_ml=25_000.0, slm_max=150_000.0,
        created_at=datetime.utcnow(),
    )
    results = run_simulation(FairInputs(), iterations=2_000, seed=11)
    return vendor, tenant, assessment, results


def test_filename_convention():
    name = report_filename("Client One", "Acme Corp!", "docx")
    assert name.startswith("VendorFAIR_ClientOne_AcmeCorp_")
    assert name.endswith(".docx")


def test_chart_renders_png(fixtures):
    _, _, _, results = fixtures
    png = render_risk_chart(results)
    assert png[:8] == b"\x89PNG\r\n\x1a\n"


def test_docx_export(tmp_path, fixtures, monkeypatch):
    monkeypatch.setattr("config.ANTHROPIC_API_KEY", "")  # force static fallbacks
    vendor, tenant, assessment, results = fixtures
    path = generate_docx(vendor, tenant, assessment, results, output_dir=tmp_path)
    doc = Document(path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    text = "\n".join(p.text for p in doc.paragraphs)
    assert any("Executive Summary" in h for h in headings)
    assert any("Reviewer Sign-off" in h for h in headings)
    assert any("Appendix" in h for h in headings)
    assert "Acme SaaS" in text
    # Chart image must be embedded
    assert any(s.startswith("image") for s in
               (r.partname.ext.lstrip(".") for r in doc.part.package.parts)
               ) or len(doc.inline_shapes) >= 1


def test_pdf_export(tmp_path, fixtures, monkeypatch):
    monkeypatch.setattr("config.ANTHROPIC_API_KEY", "")
    vendor, tenant, assessment, results = fixtures
    path = generate_pdf(vendor, tenant, assessment, results, output_dir=tmp_path)
    data = open(path, "rb").read()
    assert data[:5] == b"%PDF-"
    assert len(data) > 10_000  # chart embedded → non-trivial size
